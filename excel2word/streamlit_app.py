import streamlit as st
from pathlib import Path
import openpyxl
from openpyxl.cell.cell import Cell, MergedCell
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.shared import qn
from docx.shared import Pt
from typing import List, Tuple
import warnings
import datetime
import io
import zipfile
import tempfile
import os
import time

warnings.filterwarnings("ignore", category=UserWarning, module="openpyxl")

# ---------- CSS样式 ----------
def inject_custom_css():
    st.markdown("""
    <style>
    /* 按钮容器样式 */
    .stButton > button {
        position: relative;
        overflow: hidden;
        transition: all 0.3s ease;
    }
    
    /* 进度条样式 - 在按钮内部 */
    .progress-in-button {
        position: absolute;
        top: 0;
        left: 0;
        height: 100%;
        background: linear-gradient(90deg, rgba(255,165,0,0.3), rgba(255,165,0,0.6));
        transition: width 0.3s ease;
        z-index: 1;
    }
    
    /* 按钮文字样式 */
    .button-text {
        position: relative;
        z-index: 2;
        font-weight: bold;
    }
    
    /* 转换中按钮样式 - 橙色 */
    .converting-button {
        background-color: #f97316 !important;
        border-color: #f97316 !important;
        color: white !important;
    }
    
    .converting-button:hover {
        background-color: #ea580c !important;
        border-color: #ea580c !important;
    }
    
    /* 下载按钮样式 - 蓝色 */
    .download-button {
        background-color: #3b82f6 !important;
        border-color: #3b82f6 !important;
        color: white !important;
    }
    
    .download-button:hover {
        background-color: #2563eb !important;
        border-color: #2563eb !important;
    }
    
    /* 下载完成按钮样式 - 绿色 */
    .download-complete-button {
        background-color: #10b981 !important;
        border-color: #10b981 !important;
        color: white !important;
    }
    
    .download-complete-button:hover {
        background-color: #059669 !important;
        border-color: #059669 !important;
    }
    
    /* 窄按钮容器 */
    .narrow-button-container {
        max-width: 300px;
        margin: 0 auto;
    }
    </style>
    """, unsafe_allow_html=True)

# ---------- 边框/非空判断 ----------
def has_top_border(row: Tuple[Cell, ...]) -> bool:
    return any(c.border.top and c.border.top.style for c in row)

def non_empty_cnt(row: Tuple[Cell, ...]) -> int:
    return sum(1 for c in row if c.value is not None)

# ---------- 表格区域检测 ----------
def find_tbls(ws) -> List[Tuple[int, int]]:
    """
    返回 [(start_row, end_row), ...] 1-based
    规则：
        1. 有上边框 → 必为表格行（非空单元格数不限）。
        2. 无上边框 → 只有非空≥2 才当表格行。
        3. 表格结束：遇到既无上边框、又非空<2 的行。
    """
    tbls, in_tbl, start = [], False, None
    for idx, row in enumerate(ws.iter_rows(), 1):
        top_border = has_top_border(row)
        cnt = non_empty_cnt(row)

        if not in_tbl:                    # 当前不在表内
            if top_border or cnt >= 2:    # 有边框 或 无框但非空≥2
                in_tbl, start = True, idx
        else:                             # 已在表内
            if not top_border and cnt < 2:  # 既无框又空 → 表结束
                tbls.append((start, idx - 1))
                in_tbl = False
    if in_tbl:
        tbls.append((start, ws.max_row))
    return tbls

# ---------- 计算有效列数 ----------
def effective_cols(ws, start_row: int, end_row: int) -> int:
    """返回当前表格区域里，最右一个非空单元格所在的列号（1-based）"""
    max_col = 0
    for r in range(start_row, end_row + 1):
        row = list(ws.iter_rows(min_row=r, max_row=r))[0]
        for c in range(len(row), 0, -1):          # 从右往左找
            if row[c - 1].value is not None:
                max_col = max(max_col, c)
                break
    return max_col or 1   # 至少留 1 列

# ---------- Excel 单元格 → 字符串 ----------
def fmt_value(cell: Cell) -> str:
    """兼容 MergedCell 的取值/格式化"""
    # 0. 空值
    if cell.value is None:
        return ""

    # 1. 合并单元格只能拿到 value
    if isinstance(cell, MergedCell):
        v = cell.value
    else:
        v = cell.value   # 普通单元格

    # 2. 普通单元格精细处理
    if cell.data_type == 's':
        return cell.value or ""
    if cell.is_date:
        return cell.value.strftime('%Y年%m月%d日')
    if cell.data_type == 'n' and cell.value is not None:
        nf = cell.number_format or ''
        if '%' in nf:
            return f"{cell.value:.2%}"
        if ',' in nf or '#,#' in nf:
            return f"{cell.value:,.2f}"
        return f"{cell.value:.2f}"
    return str(cell.value) if cell.value is not None else ""

# ---------- 收集 Excel 合并单元格信息 ----------
def collect_merges(ws, tbl_start: int, tbl_end: int):
    """
    返回 [(topRow, leftCol, height, width), ...]  1-based
    只收集落在当前表格区域内的合并
    """
    rngs = []
    for m in ws.merged_cells.ranges:
        # m.min_row/max_row/min_col/max_col 都是 1-based
        if m.min_row < tbl_start or m.max_row > tbl_end:
            continue
        rngs.append((m.min_row, m.min_col,
                     m.max_row - m.min_row + 1,
                     m.max_col - m.min_col + 1))
    return rngs

# ---------- 段落样式 ----------
def set_para_format(p):
    # 段落设置
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(18)
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 字体字号设置
    run = p.runs[0] if p.runs else p.add_run()
    run.font.size = Pt(10.5)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), '宋体')

# ---------- Word 表格样式 ----------
def set_cell_format(cell, text, cell_value):
    cell.text = text
    
    # 垂直居中
    tc_pr = cell._tc.get_or_add_tcPr()
    tcVAlign = OxmlElement('w:vAlign')
    tcVAlign.set(qn('w:val'), 'center')
    tc_pr.append(tcVAlign)

    # 表格段落设置
    p = cell.paragraphs[0]
    p_format = p.paragraph_format
    p_format.space_before = Pt(5)
    p_format.space_after  = Pt(5)
    p_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p_format.line_spacing = Pt(12)

    # 表格字体字号设置
    run = p.runs[0] if p.runs else p.add_run()
    run.font.size = Pt(10.5)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), '宋体')

    # 根据单元格值类型设置对齐方式
    if isinstance(cell_value, (int, float)) and not isinstance(cell_value, bool):
        p_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

# ---------- Word 表格边框 ----------
def set_tbl_borders(tbl, thick=12, dash=6):
    rows = tbl.rows 
    if not rows:
        return

    # 横向边框
    for cell in rows[0].cells:                 
        tc_pr = cell._tc.get_or_add_tcPr()     
        tc_borders = tc_pr.first_child_found_in('w:tcBorders') 
        if tc_borders is None:                  
            tc_borders = OxmlElement('w:tcBorders')  
            tc_pr.append(tc_borders)           
        top = OxmlElement('w:top')              
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), str(thick))
        top.set(qn('w:color'), '000000')
        tc_borders.append(top)

        btm = OxmlElement('w:bottom')
        btm.set(qn('w:val'), 'dotted')
        btm.set(qn('w:sz'), str(dash))
        btm.set(qn('w:color'), '000000')
        tc_borders.append(btm)
        
    for row in rows[1:-1]:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.first_child_found_in('w:tcBorders')
            if tc_borders is None:
                tc_borders = OxmlElement('w:tcBorders')
                tc_pr.append(tc_borders)
            btm = OxmlElement('w:bottom')
            btm.set(qn('w:val'), 'dotted')
            btm.set(qn('w:sz'), str(dash))
            btm.set(qn('w:color'), '000000')
            tc_borders.append(btm)

    for cell in rows[-1].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.first_child_found_in('w:tcBorders')
        if tc_borders is None:
            tc_borders = OxmlElement('w:tcBorders')
            tc_pr.append(tc_borders)
        btm = OxmlElement('w:bottom')
        btm.set(qn('w:val'), 'single')
        btm.set(qn('w:sz'), str(thick))
        btm.set(qn('w:color'), '000000')
        tc_borders.append(btm)

    # 竖向边框
    for row in rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.first_child_found_in('w:tcBorders')
            if tc_borders is None:
                tc_borders = OxmlElement('w:tcBorders')
                tc_pr.append(tc_borders)

            if idx != len(row.cells) - 1:
                right = OxmlElement('w:right')
                right.set(qn('w:val'), 'dotted')
                right.set(qn('w:sz'), str(dash))
                right.set(qn('w:color'), '000000')
                tc_borders.append(right)

# ---------- 转换函数 ----------
def excel_to_word(excel_file, doc_stream):
    """转换单个Excel文件为Word文档"""
    try:
        wb = openpyxl.load_workbook(excel_file, data_only=True)
        ws = wb.worksheets[0]
        doc = Document()

        tbl_ranges = find_tbls(ws)
        tbl_idx = 0
        row_idx = 1

        while row_idx <= ws.max_row:
            if tbl_idx < len(tbl_ranges) and row_idx == tbl_ranges[tbl_idx][0]:
                tbl_start, tbl_end = tbl_ranges[tbl_idx]

                tbl_rows = tbl_end - tbl_start + 1
                tbl_cols = effective_cols(ws, tbl_start, tbl_end)
                tbl = doc.add_table(rows=tbl_rows, cols=tbl_cols)

                for r_offset in range(tbl_rows):
                    src_row = list(ws.iter_rows(min_row=tbl_start + r_offset,
                                                max_row=tbl_start + r_offset,
                                                values_only=False))[0]
                    dest_cells = tbl.rows[r_offset].cells
                    for c_idx in range(tbl_cols):
                        cell_value = src_row[c_idx].value
                        cell_text = fmt_value(src_row[c_idx])
                        set_cell_format(dest_cells[c_idx], cell_text, cell_value)

                for (r, c, h, w) in collect_merges(ws, tbl_start, tbl_end):
                    if c - 1 + w - 1 < tbl_cols:
                        top_left = tbl.cell(r - tbl_start, c - 1)
                        btm_right = tbl.cell(r - tbl_start + h - 1, c - 1 + w - 1)
                        top_left.merge(btm_right)

                set_tbl_borders(tbl)
                row_idx = tbl_end + 1
                tbl_idx += 1
                continue

            txt = ' '.join(fmt_value(c) for c in ws[row_idx]).strip()
            p = doc.add_paragraph(txt)
            set_para_format(p)
            row_idx += 1

        doc.save(doc_stream)
        return True, None
    except Exception as e:
        return False, str(e)

# ---------- 创建ZIP字节 ----------
def create_zip_bytes(folder_path):
    """创建ZIP文件并返回bytes"""
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for root, dirs, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                rel_path = os.path.relpath(file_path, folder_path)
                zip_file.write(file_path, rel_path)
    
    zip_buffer.seek(0)
    return zip_buffer

# ---------- 自定义按钮组件 ----------
def progress_button(label, progress=0, button_type="normal", key=None):
    """创建带进度条的按钮"""
    
    button_classes = {
        "normal": "",
        "converting": "converting-button",
        "download": "download-button",
        "complete": "download-complete-button"
    }
    
    button_class = button_classes.get(button_type, "")
    
    # 构建HTML
    html = f"""
    <div class="narrow-button-container">
        <div class="stButton">
            <button class="{button_class}" id="btn_{key}" style="width: 100%; height: 50px; position: relative;">
                <div class="progress-in-button" style="width: {progress}%;"></div>
                <span class="button-text">{label}</span>
            </button>
        </div>
    </div>
    """
    
    return st.markdown(html, unsafe_allow_html=True)

# ---------- Streamlit 界面 ----------
def main():
    st.set_page_config(
        page_title="Excel转Word工具",
        page_icon="📊",
        layout="wide"
    )
    
    # 注入CSS样式
    inject_custom_css()
    
    # 初始化会话状态
    if 'converted' not in st.session_state:
        st.session_state.converted = False
    if 'converting' not in st.session_state:
        st.session_state.converting = False
    if 'progress' not in st.session_state:
        st.session_state.progress = 0
    if 'download_data' not in st.session_state:
        st.session_state.download_data = None
    if 'download_filename' not in st.session_state:
        st.session_state.download_filename = None
    if 'success_count' not in st.session_state:
        st.session_state.success_count = 0
    if 'failed_count' not in st.session_state:
        st.session_state.failed_count = 0
    if 'failed_files' not in st.session_state:
        st.session_state.failed_files = []
    if 'is_batch' not in st.session_state:
        st.session_state.is_batch = False
    if 'prev_uploaded_files' not in st.session_state:
        st.session_state.prev_uploaded_files = None
    if 'download_clicked' not in st.session_state:
        st.session_state.download_clicked = False
    
    st.title("📊 Excel转Word文档转换工具")
    
    # 文件上传区域
    uploaded_files = st.file_uploader(
        "选择Excel文件（支持多选）",
        type=['xlsx', 'xls'],
        accept_multiple_files=True,
    )
    
    if uploaded_files:
        file_count = len(uploaded_files)
        
        # 如果上传了新文件，重置转换状态
        current_files = [f.name for f in uploaded_files]
        prev_files = st.session_state.prev_uploaded_files or []
        
        if current_files != prev_files:
            st.session_state.converted = False
            st.session_state.converting = False
            st.session_state.progress = 0
            st.session_state.download_clicked = False
            st.session_state.prev_uploaded_files = current_files
        
        # 显示文件信息（包含转换结果）
        if st.session_state.converted:
            if st.session_state.failed_count == 0:
                status_text = f"📁 已选择 **{file_count}** 个文件 | ✅ 转换成功：**{st.session_state.success_count}**"
            else:
                status_text = f"📁 已选择 **{file_count}** 个文件 | ✅ 转换成功：**{st.session_state.success_count}** | ❌ 转换失败：**{st.session_state.failed_count}**"
        else:
            status_text = f"📁 已选择 **{file_count}** 个文件"
        
        st.info(status_text)
        
        # 主按钮区域 - 居中显示
        col1, col2, col3 = st.columns([1, 2, 1])
        
        with col2:
            # 创建按钮容器
            button_container = st.empty()
            
            if not st.session_state.converted:
                # 显示转换按钮
                if st.session_state.converting:
                    # 转换中按钮（带进度条）
                    progress_button("🔄 转换中...", 
                                  st.session_state.progress, 
                                  "converting", 
                                  "converting_btn")
                else:
                    # 开始转换按钮
                    if st.button("🚀 开始转换", 
                                type="primary", 
                                use_container_width=True,
                                key="start_convert_btn"):
                        # 重置之前的结果
                        st.session_state.success_count = 0
                        st.session_state.failed_count = 0
                        st.session_state.failed_files = []
                        st.session_state.download_clicked = False
                        st.session_state.converting = True
                        st.session_state.progress = 0
                        
                        # 开始转换
                        if file_count == 1:
                            # 单文件处理
                            st.session_state.is_batch = False
                            process_single_file_with_progress(uploaded_files[0], button_container)
                        else:
                            # 多文件处理
                            st.session_state.is_batch = True
                            process_multiple_files_with_progress(uploaded_files, button_container)
            
            else:
                # 显示下载区域
                if st.session_state.download_clicked:
                    # 下载完成按钮（绿色）
                    progress_button("✅ 下载完成", 100, "complete", "download_complete_btn")
                else:
                    # 下载按钮
                    if st.session_state.is_batch:
                        button_label = "📥 下载转换结果"
                    else:
                        button_label = "📥 下载转换结果"
                    
                    # 使用download_button（蓝色按钮）
                    if st.download_button(
                        label=button_label,
                        data=st.session_state.download_data,
                        file_name=st.session_state.download_filename,
                        mime="application/zip" if st.session_state.is_batch else "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary",
                        use_container_width=True,
                        key="download_file_btn"
                    ):
                        # 设置下载已点击状态
                        st.session_state.download_clicked = True
                        st.rerun()
            
            # 显示失败文件列表（在按钮下面）
            if st.session_state.failed_files:
                with st.expander(f"📛 转换失败的文件 ({st.session_state.failed_count}个)", expanded=False):
                    for file_name, error in st.session_state.failed_files:
                        st.error(f"**{file_name}**: {error}")

def process_single_file_with_progress(uploaded_file, button_container):
    """单文件处理（带进度条）"""
    try:
        # 模拟进度更新
        st.session_state.progress = 25
        button_container.markdown(f"""
        <div class="narrow-button-container">
            <div class="stButton">
                <button class="converting-button" style="width: 100%; height: 50px; position: relative;">
                    <div class="progress-in-button" style="width: {st.session_state.progress}%;"></div>
                    <span class="button-text">🔄 转换中... 25%</span>
                </button>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        time.sleep(0.5)  # 模拟处理时间
        
        # 创建临时文件进行转换
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            st.session_state.progress = 75
            button_container.markdown(f"""
            <div class="narrow-button-container">
                <div class="stButton">
                    <button class="converting-button" style="width: 100%; height: 50px; position: relative;">
                        <div class="progress-in-button" style="width: {st.session_state.progress}%;"></div>
                        <span class="button-text">🔄 转换中... 75%</span>
                    </button>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            success, error = excel_to_word(uploaded_file, tmp_file.name)
            
            if success:
                with open(tmp_file.name, 'rb') as f:
                    doc_bytes = f.read()
                
                # 保存到会话状态
                st.session_state.download_data = doc_bytes
                st.session_state.download_filename = uploaded_file.name.replace('.xlsx', '.docx').replace('.xls', '.docx')
                st.session_state.success_count = 1
                st.session_state.converted = True
            else:
                st.session_state.failed_count = 1
                st.session_state.failed_files = [(uploaded_file.name, error)]
                st.session_state.converted = True
            
            # 清理临时文件
            os.unlink(tmp_file.name)
        
        # 完成进度
        st.session_state.progress = 100
        st.session_state.converting = False
        
    except Exception as e:
        st.session_state.failed_count = 1
        st.session_state.failed_files = [(uploaded_file.name, str(e))]
        st.session_state.converted = True
        st.session_state.converting = False

def process_multiple_files_with_progress(uploaded_files, button_container):
    """多文件处理（带进度条）"""
    # 创建临时文件夹
    with tempfile.TemporaryDirectory() as temp_dir:
        output_folder = os.path.join(temp_dir, "转换结果")
        os.makedirs(output_folder)
        
        success_count = 0
        failed_files = []
        
        # 处理每个文件
        for idx, uploaded_file in enumerate(uploaded_files):
            progress = int(((idx + 1) / len(uploaded_files)) * 100)
            st.session_state.progress = progress
            
            # 更新按钮进度
            button_container.markdown(f"""
            <div class="narrow-button-container">
                <div class="stButton">
                    <button class="converting-button" style="width: 100%; height: 50px; position: relative;">
                        <div class="progress-in-button" style="width: {progress}%;"></div>
                        <span class="button-text">🔄 转换中... {idx+1}/{len(uploaded_files)}</span>
                    </button>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            try:
                # 生成输出文件名
                doc_filename = uploaded_file.name.replace('.xlsx', '.docx').replace('.xls', '.docx')
                output_path = os.path.join(output_folder, doc_filename)
                
                # 转换文件
                success, error = excel_to_word(uploaded_file, output_path)
                
                if success:
                    success_count += 1
                else:
                    failed_files.append((uploaded_file.name, error))
                    
            except Exception as e:
                failed_files.append((uploaded_file.name, str(e)))
        
        # 保存结果到会话状态
        if success_count > 0:
            # 创建ZIP文件
            zip_buffer = create_zip_bytes(output_folder)
            
            # 保存到会话状态
            st.session_state.download_data = zip_buffer.getvalue()
            st.session_state.download_filename = f"Excel转Word_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
            st.session_state.success_count = success_count
            st.session_state.failed_count = len(failed_files)
            st.session_state.failed_files = failed_files
            st.session_state.converted = True
        else:
            # 即使全部失败也要保存状态
            st.session_state.success_count = 0
            st.session_state.failed_count = len(failed_files)
            st.session_state.failed_files = failed_files
            st.session_state.download_filename = f"Excel转Word_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
            
            # 创建一个空的ZIP文件
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                # 创建一个空文件说明
                info_file_content = "所有文件转换失败，请查看失败详情。".encode()
                zip_file.writestr("转换说明.txt", info_file_content)
            zip_buffer.seek(0)
            
            st.session_state.download_data = zip_buffer.getvalue()
            st.session_state.converted = True
        
        # 完成进度
        st.session_state.progress = 100
        st.session_state.converting = False

# ---------- 侧边栏 ----------
def sidebar_info():
    with st.sidebar:
        st.markdown("## ℹ️ 使用说明")
        st.markdown("""
        ### 简洁操作流程：
        1. **上传文件** 
        2. **点击转换** 
        3. **下载结果** 
        
        """)
        
        st.markdown("---")
        
        st.markdown("### 📊 支持格式")
        st.markdown("""
        **输入**：
        - Microsoft Excel (.xlsx)
        - Excel 97-2003 (.xls)
        
        **输出**：
        - Microsoft Word (.docx)
        """)
        
        st.markdown("---")
        
        st.markdown("### ⚠️ 注意事项")
        st.markdown("""
        1. 仅处理第一个工作表
        2. 大文件请耐心等待
        """)

if __name__ == "__main__":
    sidebar_info()
    main()
